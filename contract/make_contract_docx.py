# -*- coding: utf-8 -*-
"""Generate CONTRACT.docx for PATHMOLE Expert from the agreed terms."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEAL = RGBColor(0x0F, 0x6E, 0x6E)
DARK = RGBColor(0x1A, 0x2B, 0x33)
GRAY = RGBColor(0x55, 0x66, 0x66)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK


def heading(text, size=15, color=TEAL, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text="", bold=False, color=DARK, size=11, align=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def make_table(rows, header=True, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            pcell = cells[j].paragraphs[0]
            run = pcell.add_run(str(val))
            run.font.size = Pt(10.5)
            if header and i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif j == 0:
                run.bold = True
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t


# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("COLLABORATION AGREEMENT")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = TEAL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(10)
r = sub.add_run("Website Development & Doctor Newsletter  —  PATHMOLE Expert")
r.font.size = Pt(12.5)
r.font.color.rgb = GRAY

make_table([
    ["Client", "Dr. Arpan Gandhi & Dr. Ashok (PATHMOLE Expert)"],
    ["Developer", "Krishna Singh  —  7017737354  —  krishna191217@gmail.com"],
    ["Date", "August 2026"],
    ["Project Deadline", "August 17, 2026"],
], header=False, widths=[1.8, 4.7])

# ---------- Overview ----------
heading("Overview")
para("This agreement outlines the terms for building a new, modern website for PATHMOLE "
     "Expert — built from scratch in HTML, CSS, and JavaScript — with a rule-based chatbot "
     "and a doctor case-study newsletter that publishes on the website and goes out by "
     "email to your referring doctors.")
para("What you get:", bold=True, space_after=2)
bullet(" — Clean, clinical design that reflects the lab's credibility", "Looks premium")
bullet(" — 2–3x faster than WordPress, better for user engagement", "Performs faster")
bullet(" — Accreditation, publications, and case studies showcased properly", "Builds trust")
bullet(" — De-identified case studies emailed to your referring doctors", "Keeps doctors engaged")
bullet(" — SEO-friendly structure for better Google visibility", "Grows organically")
bullet(" — No WordPress vulnerabilities, plugin breaks, or database issues", "Stays secure")

# ---------- Why static ----------
heading("Why Static HTML Over WordPress?")
make_table([
    ["Factor", "WordPress", "HTML / CSS / JS"],
    ["Speed", "Slower", "2–3x faster"],
    ["Security", "Vulnerable", "Highly secure"],
    ["Maintenance", "Updates needed", "Minimal"],
    ["Reliability", "Can break", "Never breaks"],
], widths=[2.0, 2.2, 2.3])

# ---------- Scope ----------
heading("Scope of Work")
para("Core pages", bold=True, space_after=2)
bullet("Home — sticky top bar (Phone, Hours, Call Now, WhatsApp), hero, tagline, mission, enquiry CTA")
bullet("About Us / About the Lab — founders, partnership, vision, quality philosophy")
bullet("Services — Histopathology & Molecular Diagnostics (histopathology, cytology, IHC, molecular panels)")
bullet("Quality & Accreditation — NABL/CAP goals, SOPs, turnaround-time commitment")
para("Tests & reporting", bold=True, space_after=2)
bullet("Test List — tests presented with categories, symptoms/indications, and details; guided by the chatbot 'find a test' flow (no pricing shown — pricing handled via enquiry/contact)")
bullet("Doctor & Patient Login — redirects to your existing Reporting Portal (link only)")
para("Content & media", bold=True, space_after=2)
bullet("Research & Publications")
bullet("Case Studies & Newsletter — de-identified case studies published on the website and emailed to referring doctors")
bullet("Gallery — facility photos plus equipment/machines, their uses, and related information")
bullet("Videos — embedded YouTube section")
para("Patients & engagement", bold=True, space_after=2)
bullet("Patients — information plus a downloadable Patient Form (PDF); no online data storage")
bullet("FAQ, Testimonials, Careers")
para("Contact & extras", bold=True, space_after=2)
bullet("Contact Us — Google Maps, enquiry form, address, details")
bullet("Social Media Integration")
bullet("Custom Rule-Based Chatbot — answers common questions, guides users to WhatsApp / phone")

# ---------- Phases ----------
heading("Collaboration Phases")
heading("Phase 1: Website Development  —  ₹26,000", size=12.5, color=DARK, space_before=8, space_after=2)
para("Timeline: August 2026 – August 17, 2026", color=GRAY, space_after=4)
para("Key Milestones:", bold=True, space_after=2)
bullet("Project Start: on advance payment (received)")
bullet("First Draft Delivery: August 11–12, 2026")
bullet("Deployment / Launch: August 17, 2026")
para("Deliverables:", bold=True, space_after=2)
bullet("Complete website built in HTML / CSS / JS from scratch (all sections above)")
bullet("Modern, fast, premium design reflecting the lab's brand")
bullet("Rule-based chatbot")
bullet("Doctor case-study newsletter setup (professional email + newsletter tool, doctor list import, branded template)")
bullet("Basic testing and quality check")
bullet("Published on the lab's domain")
para("Payment: ₹7,000 advance to begin (received)  ·  ₹19,000 on completion, before go-live",
     bold=True, space_after=6)

heading("Phase 2: Monthly Maintenance  —  ₹2,000 / month", size=12.5, color=DARK, space_before=8, space_after=2)
para("Ongoing · billed monthly (pay at end of month) · cancellable with 30 days' notice",
     color=GRAY, space_after=4)
para("Included:", bold=True, space_after=2)
bullet("Text and image updates")
bullet("Minor design tweaks and bug fixes")
bullet("Mobile responsive adjustments")
bullet("Contact / social link updates")
bullet("Publishing new case studies and sending the newsletter")
para("Not Included (billed separately):", bold=True, space_after=2)
bullet("New pages or major redesign")
bullet("New features (booking, payment, etc.)")
bullet("Backend / online patient-record storage")
bullet("Reporting Portal API integration (beyond the redirect link)")
bullet("Complex animations / video production")
bullet("Content writing, logo design, photography")

# ---------- Investment ----------
heading("Investment Summary")
make_table([
    ["Phase", "Description", "Price"],
    ["Phase 1", "Website + Chatbot + Newsletter setup", "₹26,000"],
    ["Total", "", "₹26,000"],
    ["Phase 2", "Monthly Maintenance (optional, ongoing)", "₹2,000 / month"],
], widths=[1.3, 4.2, 1.5])
para("Overall Project Timeline: August 2026 – August 17, 2026", color=GRAY, space_after=4)

# ---------- Timeline ----------
heading("Project Timeline")
make_table([
    ["Milestone", "Date"],
    ["Project Start", "On advance (received)"],
    ["First Draft", "August 11–12, 2026"],
    ["Deployment / Launch", "August 17, 2026"],
], widths=[3.2, 3.3])

# ---------- Payment ----------
heading("Payment Schedule")
make_table([
    ["Milestone", "Amount"],
    ["Advance (received)", "₹7,000"],
    ["On completion, before go-live", "₹19,000"],
    ["Monthly Maintenance (optional)", "₹2,000 / month"],
], widths=[3.8, 2.7])
para("Payment method: [ UPI / bank details ]", color=GRAY)

# ---------- What I need ----------
heading("What I Need From You")
para("To meet the August 17 launch, please share promptly:", space_after=2)
numbered("Domain — purchased by you (used for the website and email; ~₹500–1,300/year). Email + newsletter tools run on free plans.")
numbered("Logo, brand colors, and any brand guidelines")
numbered("Test list data (test names, categories, symptoms/indications, and any sample/prep info) — no pricing needed on the site")
numbered("Content for About, Services, Quality, Publications")
numbered("Photos for Gallery (facility + equipment) and YouTube video links")
numbered("Referring-doctor list (name + email) for the newsletter")
numbered("The downloadable Patient Form (PDF)")
numbered("Contact details, working hours, social media links")

# ---------- Terms ----------
heading("Terms & Notes")
bullet(" Up to 2 rounds of revisions within scope are included; extra revisions or new scope are billed separately.", "Revisions:")
bullet(" Bugs in the delivered scope are fixed free for 14 days after launch. After that, a ₹1,000 charge applies, and continued support is covered by the agreed ₹2,000/month maintenance.", "Support:")
bullet(" On full payment of the balance, ownership of the website transfers to you. All accounts and the domain remain yours; I set them up on your behalf. I may show the completed site in my portfolio unless you object in writing.", "Ownership:")
bullet(" All case studies — on the website and in email — are fully de-identified (no patient name, ID, photo, or identifying detail) with a disclaimer, and the newsletter includes a working unsubscribe. You confirm de-identification and hold any required consent. The website does not collect or store patient records. Your doctor list stays confidential and is used only for this project.", "Privacy (DPDP Act, 2023):")
bullet(" Depends on timely delivery of your content and approvals; content delays extend the date accordingly.", "Timeline:")
bullet(" I work as an independent contractor. This is the full agreement between us; changes must be agreed in writing. Governed by the laws of India. The advance is non-refundable.", "General:")

# ---------- Next steps ----------
heading("Next Steps")
numbered("Confirm acceptance of these terms")
numbered("Share the content listed above so we stay on schedule")
numbered("Weekly progress updates shared throughout")

para("")
para("Thank you, Dr. Arpan and Dr. Ashok. Let's build a website that reflects PATHMOLE "
     "Expert's expertise and keeps your referring doctors engaged.", color=GRAY)

# ---------- Signatures ----------
heading("Accepted & Agreed by")
para("")
para("Client Signature: __________________________     Date: ______________")
para("Name: Dr. Arpan Gandhi", space_after=10)
para("Client Signature: __________________________     Date: ______________")
para("Name: Dr. Ashok", space_after=10)
para("Developer Signature: _______________________     Date: ______________")
para("Name: Krishna Singh  —  7017737354")

out = r"D:\Freelance\PATHMOLE Website proejct\contract\PATHMOLE-Expert-Contract.docx"
doc.save(out)
print("Saved:", out)
