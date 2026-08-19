# -*- coding: utf-8 -*-
"""
Generates 'PATHMOLE-Content-Request.docx' — a structured content-collection
template the client fills in, page by page, so every placeholder on the live
site can be replaced with real content.

Run:  python scripts/make_content_request_docx.py
Needs: pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x23, 0x2C, 0x8E)
MAGENTA = RGBColor(0xEC, 0x00, 0x8C)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def _shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)

def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)

def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)

def page_heading(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num}.  {text}")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = NAVY

def page_file(fname, purpose):
    p = doc.add_paragraph()
    r = p.add_run(f"File: {fname}   |   {purpose}")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MAGENTA
    p.paragraph_format.space_after = Pt(6)

def section_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x22, 0x70)

def fixed_line(label, value):
    """A heading/label that is already set on the site (client just confirms)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY

def prompt(question, hint=None, lines=1):
    """A thing the client must supply. Renders a labelled fill-in box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("▸ " + question)
    r.font.bold = True
    r.font.size = Pt(10.5)
    if hint:
        ph = doc.add_paragraph()
        ph.paragraph_format.space_after = Pt(2)
        rh = ph.add_run(hint)
        rh.font.size = Pt(9)
        rh.font.italic = True
        rh.font.color.rgb = GREY
    # answer box (single-cell table)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _shade(cell, "F7F8FC")
    cell.paragraphs[0].add_run("\n" * lines)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ============================ COVER ============================
title("PATHMOLE Expert Lab — Website Content Request")
subtitle("Please fill in the boxes below. Every item corresponds to a real section on your website. "
         "Leave a box blank only if you want us to keep the current draft wording. "
         "Return this document (or type answers inline) and we will load it into the site.")

hr()
section_label("How to use this form")
for line in [
    "Blue headings = section on the site. The wording after “Fixed:” is already on the page — tell us only if you want it changed.",
    "▸ marked items = content we need FROM YOU. Type into the grey box under each.",
    "Files/photos: don’t paste into this doc — send them in a folder named the same as the section (e.g. “Gallery”, “Patient Form”).",
    "Two hard rules we keep on the site: (1) NO pricing shown anywhere — price enquiries are routed to phone/WhatsApp; (2) All case studies are fully de-identified (no patient name, ID, photo, or identifying detail).",
]:
    b = doc.add_paragraph(style="List Bullet")
    b.add_run(line).font.size = Pt(10)

section_label("Global items (used across the whole site)")
prompt("Confirm the lab phone number, WhatsApp number and email.",
       "Currently on site: +91 98998 22375 (call & WhatsApp), pathmolelab@gmail.com")
prompt("Confirm the full address and opening hours.",
       "On site: Building No. 1164/1, 1st Floor, Shri JP Tower, New Railway Road, Opp. Fire Station, "
       "Dayanand Colony, Sector 6, Gurugram (Haryana). Open daily 8:00 AM – 8:00 PM.")
prompt("Social media links (Facebook / Instagram / LinkedIn / other).", "Leave blank any you don’t have.")
prompt("The Reports-Login / Reporting-Portal web address.", "The “Reports Login” button links here.")
prompt("Vector logo file (SVG / AI / PDF) and website domain name.",
       "Send the logo as a file; type the domain here if decided.")

doc.add_page_break()

# ============================ PAGES ============================

# 1 HOME
page_heading(1, "Home page")
page_file("index.html", "The landing page")
section_label("Hero (top banner)")
fixed_line("Fixed headline", "“Precision diagnostics your clinicians can trust”")
fixed_line("Fixed tagline", "“Precision in Diagnosis. Confidence in Care.”")
prompt("Change the headline or tagline? (optional)", lines=2)
section_label("Announcement strip")
prompt("Latest announcement or news headline to show at the top.",
       "e.g. a new test launched, an award, a new case study. Optional — we can hide the strip.")
section_label("Trust numbers (the four stats)")
prompt("Years of expertise / Referring clinicians / Tests offered — give real figures.",
       "On site these show as [XX]+. Also list any accreditation (e.g. NABL) ONLY if confirmed.")
section_label("Testimonials (3 quotes from referring doctors)")
prompt("Testimonial 1 — quote + doctor name + speciality/city.", lines=2)
prompt("Testimonial 2 — quote + doctor name + speciality/city.", lines=2)
prompt("Testimonial 3 — quote + doctor name + speciality/city.", lines=2)

# 2 ABOUT
page_heading(2, "About Us")
page_file("about.html", "Who you are and your story")
prompt("The lab’s story / philosophy / what makes it distinctive.",
       "A few sentences to a couple of paragraphs.", lines=4)
section_label("Leadership")
prompt("Dr. Arpan Gandhi — full title, qualifications, short bio.", lines=3)
prompt("Second principal — CONFIRM name & title (contract says “Dr. Ashok”, site draft says “Mr. Ashok Yadav”) + qualifications + short bio.", lines=3)
prompt("Add any other team members here (name, title, bio).", lines=2)

# 3 SERVICES
page_heading(3, "Services")
page_file("services.html", "What the lab does")
fixed_line("Sections on page", "Histopathology  ·  Molecular Diagnostics  ·  Immunohistochemistry (IHC)  ·  Quality & Turnaround")
prompt("For EACH of the four services — a short description you’re happy with (or approve the current draft).", lines=4)
prompt("Accreditation details to add under Quality (only if confirmed).")

# 4 TESTS
page_heading(4, "Test List")
page_file("tests.html + data/tests.js", "The searchable list of tests offered")
prompt("The FULL, confirmed list of tests. For each: test name, category (Histopathology / Molecular / IHC / other), and a one-line description.",
       "The site currently lists 15 sample tests. Reminder: NO prices — pricing is handled by phone/WhatsApp.", lines=6)

# 5 CASE STUDIES
page_heading(5, "Case Studies")
page_file("case-studies/index.html", "De-identified teaching cases")
prompt("Case studies to publish. For each: a title, the teaching point/summary, discipline (Histo/Molecular/IHC), and month/year.",
       "MUST be fully de-identified — no patient name, ID, photo, or identifying detail. Send any images separately.", lines=5)

# 6 PUBLICATIONS
page_heading(6, "Research & References (Publications)")
page_file("publications.html", "Papers and reference guidelines")
prompt("The lab’s own publications / posters / presentations (title, authors, where published, year, link).", lines=3)
fixed_line("Already listed (reference guidelines)", "WHO Classification of Tumours 5th ed.; ASCO–CAP HER2 (2023); CAP–IASLC–AMP molecular; CAP–AMP MMR/MSI (2022)")
prompt("Keep, remove, or add to the reference-guideline list above?")

# 7 QUALITY
page_heading(7, "Quality & Accreditation")
page_file("quality.html", "Quality systems & accreditations")
prompt("Describe your quality process (SOPs, internal QC, expert sign-out).", lines=3)
prompt("List accreditations to display (NABL / CAP / other) — ONLY those actually held.", "Include certificate numbers if you want them shown.")

# 8 PHYSICIANS
page_heading(8, "Physicians & Team")
page_file("physicians.html", "The people behind the reports")
prompt("Same team details as About (name, title, qualifications, bio, and a photo if you’d like).",
       "Send photos as image files named per person.", lines=3)

# 9 GALLERY
page_heading(9, "Gallery (Lab photos)")
page_file("gallery.html", "Photos of the lab, equipment, team")
prompt("Send 6+ photos of the lab / equipment / premises / team, with a one-line caption for each.",
       "Send as image files (JPG/PNG) in a folder “Gallery” — not pasted into this document.", lines=2)

# 10 VIDEOS
page_heading(10, "Videos")
page_file("videos.html", "Explainer / lab-tour videos")
prompt("Any videos to embed — give the YouTube/Vimeo link and a title for each.",
       "Or send the video files if not yet uploaded. Optional.", lines=2)

# 11 PATIENTS + FORM
page_heading(11, "For Patients  +  Patient Form download")
page_file("patients.html", "Patient info and the downloadable form")
prompt("Patient information text — what to expect, sample collection, timings, how to collect reports.", lines=4)
prompt("The Patient Form itself — send the final PDF (or the content, and we’ll lay it out).",
       "It will download from assets/pathmole-patient-form.pdf. NOTE: the form is download-only — the site does not store any patient data.")

# 12 FAQ
page_heading(12, "FAQ")
page_file("faq.html", "Frequently asked questions")
fixed_line("Questions already answered", "Timings · Location · How to get pricing · Are case studies identifiable · How to access reports")
prompt("Approve/correct the current answers, and add any more Q&As you get asked often.", lines=3)

# 13 CAREERS
page_heading(13, "Careers")
page_file("careers.html", "Job openings")
prompt("Current openings — for each: role title, responsibilities, requirements. Confirm the applications email.",
       "On site, applications go to pathmolelab@gmail.com. Leave blank if no openings.", lines=3)

# 14 CONTACT
page_heading(14, "Contact Us")
page_file("contact.html", "Contact details + enquiry form")
fixed_line("Enquiry form", "Already built (Name, Clinic/Hospital, Phone, Email, Message). It emails the lab; no data stored on the site.")
prompt("Which email should enquiry-form submissions be sent to?", "Also: do you have a form-service account (e.g. Formspree), or should we set one up?")
prompt("Google Maps location link to embed on the page.", "Share the Google Maps ‘share’ or ‘embed’ link for the Sector 6 address.")

# ---- closing ----
doc.add_page_break()
section_label("Anything else")
prompt("Anything not covered above that you’d like on the website?", lines=4)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
r = p.add_run("Thank you — once we receive this, we’ll replace every placeholder and share a preview before launch.")
r.font.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

out = "docs/PATHMOLE-Content-Request.docx"
doc.save(out)
print("Wrote", out)
